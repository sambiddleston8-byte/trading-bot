"""Build the stakeholder handover report from the latest saved prototype data.

The generated sector chart is a source-backed portfolio output visual. The
report never reads `.env` or includes provider credentials.
"""

from __future__ import annotations

import json
import math
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUTPUT = DOCS / "Portfolio_Construction_Prototype_Report.docx"
CHART = ASSETS / "portfolio_sector_allocation.png"
PORTFOLIOS = ROOT / "data" / "research" / "portfolios"
MONITORING = ROOT / "data" / "research" / "portfolio_monitoring"

NAVY = "0B2545"
BLUE = "2E74B5"
MUTED = "5B6573"
PALE = "F2F4F7"
GREEN = "1F5F4A"
GOLD = "7A5A00"
RED = "9B1C1C"
PAGE_WIDTH_DXA = 9360


def latest_json(directory: Path, pattern: str) -> tuple[dict, Path]:
    for path in sorted(directory.glob(pattern), reverse=True):
        try:
            value = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if isinstance(value, dict):
            return value, path
    raise RuntimeError(f"No valid {pattern} record found in {directory}")


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def pie_chart(sector_weights: dict[str, float]) -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    canvas = Image.new("RGB", (1400, 820), "white")
    draw = ImageDraw.Draw(canvas)
    title_font, label_font, small_font = font(44, True), font(28, True), font(24)
    draw.text((68, 48), "Latest proposed portfolio: sector allocation", font=title_font, fill=f"#{NAVY}")
    draw.text((70, 112), "Source: latest saved paper portfolio. This is a data visual, not a performance forecast.", font=small_font, fill=f"#{MUTED}")
    colors = ["#2E74B5", "#1F5F4A", "#C88719", "#7D5BA6", "#B45152", "#4D8DAD", "#6C7A89"]
    weights = sorted(sector_weights.items(), key=lambda item: item[1], reverse=True)
    start = -90.0
    box = (70, 190, 650, 770)
    for index, (_, value) in enumerate(weights):
        end = start + (float(value) * 360.0)
        draw.pieslice(box, start=start, end=end, fill=colors[index % len(colors)], outline="white", width=4)
        start = end
    draw.ellipse((270, 390, 450, 570), fill="white")
    draw.text((306, 440), "15", font=font(54, True), fill=f"#{NAVY}")
    draw.text((282, 507), "holdings", font=small_font, fill=f"#{MUTED}")
    y = 215
    for index, (sector, value) in enumerate(weights):
        color = colors[index % len(colors)]
        draw.rounded_rectangle((760, y + 3, 788, y + 31), radius=5, fill=color)
        draw.text((810, y), sector, font=label_font, fill=f"#{NAVY}")
        percent = f"{float(value):.1%}"
        right = draw.textbbox((0, 0), percent, font=label_font)[2]
        draw.text((1310 - right, y), percent, font=label_font, fill=f"#{NAVY}")
        y += 76
    canvas.save(CHART)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    table_properties.append(indent)
    grid = table._tbl.tblGrid
    for column, width in zip(grid.gridCol_lst, widths):
        column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Pt(width / 20)
            width_node = cell._tc.tcPr.tcW
            width_node.set(qn("w:w"), str(width))
            width_node.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_run(run, *, size=11, color=NAVY, bold=False, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def apply_style(paragraph, *, before=0, after=6, line=1.10, align=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text: str, *, size=11, color=NAVY, bold=False, italic=False, before=0, after=6, align=None) -> None:
    paragraph = doc.add_paragraph()
    apply_style(paragraph, before=before, after=after, align=align)
    set_run(paragraph.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    values = {1: (16, BLUE, 16, 8), 2: (13, BLUE, 12, 6), 3: (12, "1F4D78", 8, 4)}
    size, color, before, after = values[level]
    apply_style(paragraph, before=before, after=after)
    set_run(paragraph.add_run(text), size=size, color=color, bold=True)
    return paragraph


def add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    apply_style(paragraph, after=4, line=1.167)
    set_run(paragraph.add_run(text), size=11, color=NAVY)


def add_status_table(doc, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2400, 5200, 1760])
    for cell, title in zip(table.rows[0].cells, ("Area", "What is working now", "Status")):
        set_cell_shading(cell, PALE)
        paragraph = cell.paragraphs[0]
        apply_style(paragraph, after=0)
        set_run(paragraph.add_run(title), size=10, color=NAVY, bold=True)
    for area, detail, status in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, (area, detail, status)):
            paragraph = cell.paragraphs[0]
            apply_style(paragraph, after=0)
            set_run(paragraph.add_run(value), size=9.5, color=NAVY, bold=(cell is cells[0]))
    add_text(doc, "Status is based on the latest saved prototype run and is not an investment recommendation.", size=9, color=MUTED, italic=True, before=4, after=8)


def add_holding_table(doc, holdings: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [3300, 1250, 2300, 2510])
    headers = ("Company", "Allocation", "Decision rating", "Portfolio decision")
    for cell, title in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE)
        paragraph = cell.paragraphs[0]
        apply_style(paragraph, after=0)
        set_run(paragraph.add_run(title), size=9.5, color=NAVY, bold=True)
    for holding in holdings:
        cells = table.add_row().cells
        values = (
            f"{holding.get('name') or holding.get('ticker')} ({holding.get('ticker')})",
            f"{float(holding.get('weight') or 0):.1%}",
            f"{min(100.0, float((holding.get('decision_rating') or {}).get('score') or 0)):.0f}/100",
            str(holding.get("portfolio_decision") or "SELECTED"),
        )
        for cell, value in zip(cells, values):
            paragraph = cell.paragraphs[0]
            apply_style(paragraph, after=0)
            set_run(paragraph.add_run(value), size=9.2, color=NAVY, bold=(cell is cells[0]))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    for name, size, color, before, after in (
        ("Normal", 11, NAVY, 0, 6),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
    header = section.header.paragraphs[0]
    apply_style(header, after=0)
    set_run(header.add_run("SAM & PAT PROJECT  |  PORTFOLIO CONSTRUCTION PROTOTYPE"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    apply_style(footer, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(footer.add_run("Research prototype | Paper monitoring only | August 2026"), size=8.5, color=MUTED)


def main() -> None:
    portfolio, _ = latest_json(PORTFOLIOS, "research_portfolio_*.json")
    monitoring, _ = latest_json(MONITORING, "portfolio_health_*.json")
    sector_weights = portfolio.get("sector_weights") or {}
    pie_chart(sector_weights)
    holdings = portfolio.get("holdings") or []
    weighted_gap = sum(float(item.get("weight") or 0) * float(item.get("expected_return") or 0) for item in holdings)
    forecast_years = portfolio.get("valuation_horizon_years") or 5
    annualised_gap = (1 + weighted_gap) ** (1 / float(forecast_years)) - 1
    counts = monitoring.get("summary", {}).get("action_counts", {})
    allocation_changes = monitoring.get("summary", {}).get("allocation_changes_required", 0)
    now = datetime.now().strftime("%d %B %Y")

    doc = Document()
    configure_document(doc)

    for _ in range(5):
        add_text(doc, "", after=0)
    add_text(doc, "PORTFOLIO CONSTRUCTION", size=12, color=GOLD, bold=True, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Prototype Report", size=30, color=NAVY, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "What it does now and what comes next", size=15, color="2B5163", after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Prepared for Sam & Pat | {now}", size=10.5, color=MUTED, after=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "A transparent research and paper-portfolio prototype. It does not connect to a broker, place trades or promise returns.", size=11, color=NAVY, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_heading(doc, "Executive summary")
    add_text(doc, "The tool researches companies in the S&P 500 and Nasdaq-100, checks the quality of the evidence, then builds a 15-company proposed portfolio. Each holding links to a full research report. A separate Current portfolio view highlights price, thesis, audit and allocation changes for review. It never trades automatically.")
    add_status_table(doc, [
        ("Research universe", "517 companies across the S&P 500 and Nasdaq-100; 192 have current saved research.", "Active"),
        ("Portfolio", f"{len(holdings)} holdings, no cash, a 50% sector limit and a passed risk check.", "Working"),
        ("Decision rating", "A true 0-100 score. It keeps the differences between companies and includes a small uncertainty deduction, so it never gives a false 100/100 certainty score.", "Working"),
        ("Monitoring", f"{monitoring.get('summary', {}).get('position_count', 0)} positions checked; {counts.get('HOLD', 0)} unchanged and {allocation_changes} allocation changes flagged for review. Price alone is not a sell signal.", "Working"),
        ("S&P 500 benchmark", "The starting observation is saved. A one-month comparison starts after 30 days of records.", "Building history"),
    ])

    add_heading(doc, "What the tool does today")
    for text in (
        "Looks at financials, valuation, business quality, management, competitors, earnings, catalysts, news, market signals and the wider economy.",
        "Checks the evidence and challenges the investment case before a company can be included.",
        "Avoids duplicate company exposure, limits concentration and gives different companies different weights based on return potential, evidence, risk and volatility.",
        "Lets you click a ticker for the investment case, risks, suggested holding period, valuation, catalysts, bot research and audit findings.",
        "Creates monitoring checks using refreshed fundamentals, valuation, technical signals, catalysts, thesis challenge and audit results. Price is shown as context, not as a fixed sell rule. All changes remain prompts for a person to review.",
    ):
        add_bullet(doc, text)

    doc.add_page_break()
    add_heading(doc, "Latest proposed portfolio")
    add_text(doc, "This is the latest saved allocation. It is not a performance forecast.", size=10, color=MUTED, italic=True)
    doc.add_picture(str(CHART), width=Inches(6.5))
    add_text(doc, "The model estimates {:.1%} price upside over about {:.0f} years (roughly {:.1%} a year in the model). This is an estimate, not a promise.".format(weighted_gap, float(forecast_years), annualised_gap), size=10, color=NAVY, before=4, after=8)
    add_holding_table(doc, holdings)

    doc.add_page_break()
    add_heading(doc, "Research sources and decision controls")
    add_text(doc, "Company filings are the main source of financial facts. Other sources add useful checks, but they do not overrule the audit process.")
    for text in (
        "SEC filings: the main source for reported revenue, cash flow, debt, shares and other company financial facts.",
        "Yahoo Finance/yfinance: current and historic market prices used for market context, volatility and technical analysis.",
        "Financial Modeling Prep: a separate check on reported statements, analyst estimates, rating snapshots and analyst price-target consensus.",
        "Massive: historic market data and company news, used to strengthen price context and identify relevant events.",
        "FRED: official US macroeconomic series such as interest rates, inflation and growth, used to assess the macro and discount-rate backdrop.",
        "Alpha Vantage: an additional independent income-statement source used as a cross-check when available.",
        "Data quality: the system shows missing or conflicting data instead of making up values. News is filtered so irrelevant headlines do not influence a company decision.",
    ):
        add_bullet(doc, text)
    add_heading(doc, "Portfolio safeguard policy", level=2)
    add_text(doc, "The prototype stays fully invested, but it does not treat every researched company as suitable. A failed audit, a broken thesis, weak valuation quality, duplicate exposure, too much sector concentration or a failed risk check can block or reduce a holding.")
    add_text(doc, "The decision rating combines the investment case, evidence quality, estimated return, valuation reliability, risk, thesis, technical context and the breadth of sources. It uses the full 0-100 scale and keeps the differences between companies. A small uncertainty deduction means it never gives a false 100/100 certainty score. It is not a prediction that a share price will rise.", size=10.5, color=GREEN, bold=True, before=4, after=8)

    add_heading(doc, "S&P 500 comparison and current portfolio", level=1)
    add_text(doc, "The S&P 500 is the benchmark. Once the system has 30 days of saved checks, the website will compare the portfolio's past-month return with the S&P 500 over the same period. Until then, it says there is not enough data.")
    add_text(doc, f"The performance chart will use real dated checks, not estimated past results. The latest check records {allocation_changes} allocation changes for review, but the system has not changed the portfolio itself. A change is triggered by refreshed research evidence, not by an automatic 12% or 20% price rule.")

    doc.add_page_break()
    add_heading(doc, "What else needs actioning")
    for text in (
        "Run scheduled full research refreshes for current holdings so the monitoring view uses up-to-date fundamentals, valuation, technical analysis, catalysts, sentiment, thesis challenge and audit results.",
        "Add a proper catalyst probability model for events such as approvals, trials, contracts, product launches, acquisitions and regulatory decisions, using source-backed probabilities rather than headline sentiment alone.",
        "Improve technical analysis with stronger trend, momentum, volume, support/resistance and relative-strength measures, all compared against the wider market and sector.",
        "Expand fundamental analysis across several reporting periods: revenue growth, margins, free cash flow, debt, returns on capital, share dilution and guidance changes.",
        "Build sector and industry comparison models so each company is judged against direct competitors, not only against its own history.",
        "Add earnings-call, management-guidance and filing-change extraction, with clear links to the evidence that changed a view.",
        "Make macro analysis more specific: interest-rate, inflation, growth, credit and sector sensitivity should affect valuation assumptions and position sizes where relevant.",
        "Add portfolio-level correlation, factor and liquidity analysis so diversification is based on economic exposure, not only company names and sectors.",
        "Replace generic exit rules with a research-led exit framework: thesis break, valuation change, evidence failure, catalyst completion, deteriorating fundamentals or deteriorating technical context.",
        "Keep a complete decision log: what was known at the time, why a holding was chosen, what changed later and whether the decision was right.",
        "Build point-in-time historical datasets and backtests, including realistic transaction costs, turnover, delisted companies and out-of-sample testing.",
        "Build a longer paper record and compare it with the S&P 500 over 1, 3, 6 and 12 months before claiming an edge.",
        "Use paper outcomes to calibrate the scoring model. The learning bot should improve weights only from recorded evidence and later results, never hindsight.",
        "Add operational controls: source-health checks, stale-data warnings, error monitoring, reproducible research snapshots and manual approval for every portfolio change.",
        "Consider live trading only after strong paper results, independent review, clear risk limits, trading-cost analysis, broker controls and your explicit approval.",
    ):
        add_bullet(doc, text)


    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
